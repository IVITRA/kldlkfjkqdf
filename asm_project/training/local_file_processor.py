"""
Smart Document Parser - Local File Processing Pipeline
Reads any file: PDF, TXT, DOCX, CSV, even scanned images with OCR
"""

import fitz  # PyMuPDF for PDF
import docx
import pandas as pd
from pathlib import Path
import re
from typing import List, Dict, Optional
import pytesseract  # OCR for images
from PIL import Image
import io
import json


class LocalDocumentProcessor:
    """
    Processes local documents and extracts structured content
    Supports: PDF, TXT, DOCX, MD, CSV, JSON
    """
    
    def __init__(self, ocr_enabled: bool = True, language: str = 'ara+eng'):
        self.ocr_enabled = ocr_enabled
        self.language = language
        self.supported_formats = ['.pdf', '.txt', '.docx', '.md', '.csv', '.json']
        
    def process_file(self, file_path: Path) -> Dict:
        """
        Processes a single file and returns structured content
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._process_pdf(file_path)
        elif suffix in ['.txt', '.md']:
            return self._process_text(file_path)
        elif suffix == '.docx':
            return self._process_docx(file_path)
        elif suffix == '.csv':
            return self._process_csv(file_path)
        elif suffix == '.json':
            return self._process_json(file_path)
        else:
            raise ValueError(f"Unsupported format: {suffix}")
    
    def process_directory(self, dir_path: Path, recursive: bool = True) -> List[Dict]:
        """
        Processes all supported files in a directory
        """
        dir_path = Path(dir_path)
        documents = []
        
        pattern = '**/*' if recursive else '*'
        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    doc = self.process_file(file_path)
                    documents.append(doc)
                    print(f"✓ Processed: {file_path.name}")
                except Exception as e:
                    print(f"✗ Error processing {file_path}: {e}")
                    
        return documents
    
    def _process_pdf(self, path: Path) -> Dict:
        """
        Extracts text from PDF with layout preservation
        Uses OCR for scanned documents
        """
        doc = fitz.open(path)
        pages = []
        
        for page_num, page in enumerate(doc):
            # Direct text extraction
            text = page.get_text()
            
            # If text is minimal (scanned image), use OCR
            if len(text.strip()) < 100 and self.ocr_enabled:
                print(f"  → Using OCR for page {page_num + 1}")
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better accuracy
                img = Image.open(io.BytesIO(pix.tobytes()))
                text = pytesseract.image_to_string(img, lang=self.language)
            
            pages.append({
                'page': page_num + 1,
                'text': text,
                'word_count': len(text.split()),
                'images': []  # Can extract images too
            })
        
        return {
            'source': str(path),
            'type': 'pdf',
            'pages': pages,
            'total_pages': len(pages),
            'total_words': sum(p['word_count'] for p in pages),
            'metadata': doc.metadata
        }
    
    def _process_docx(self, path: Path) -> Dict:
        """
        Extracts text with heading and list structure preservation
        """
        doc = docx.Document(path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine paragraph type (heading, normal text, list)
                style = para.style.name if para.style else 'Normal'
                level = self._get_heading_level(style)
                
                paragraphs.append({
                    'text': para.text,
                    'style': style,
                    'level': level,
                    'is_heading': level > 0
                })
        
        # Extract tables if present
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'source': str(path),
            'type': 'docx',
            'paragraphs': paragraphs,
            'tables': tables,
            'total_paragraphs': len(paragraphs),
            'total_words': sum(len(p['text'].split()) for p in paragraphs)
        }
    
    def _process_text(self, path: Path) -> Dict:
        """
        Processes plain text or markdown files
        """
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Detect headings in markdown
        headings = []
        if path.suffix.lower() == '.md':
            heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
            for match in heading_pattern.finditer(text):
                headings.append({
                    'level': len(match.group(1)),
                    'text': match.group(2)
                })
        
        return {
            'source': str(path),
            'type': 'text',
            'text': text,
            'headings': headings,
            'word_count': len(text.split()),
            'line_count': text.count('\n') + 1
        }
    
    def _process_csv(self, path: Path) -> Dict:
        """
        Processes CSV files and converts to structured text
        """
        df = pd.read_csv(path)
        
        # Convert to text representation
        text_rows = []
        for idx, row in df.iterrows():
            row_text = ', '.join([f"{col}: {val}" for col, val in row.items()])
            text_rows.append(row_text)
        
        return {
            'source': str(path),
            'type': 'csv',
            'columns': list(df.columns),
            'num_rows': len(df),
            'text': '\n'.join(text_rows),
            'preview': df.head(10).to_dict()
        }
    
    def _process_json(self, path: Path) -> Dict:
        """
        Processes JSON files
        """
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert to text
        if isinstance(data, dict):
            text = json.dumps(data, ensure_ascii=False, indent=2)
        elif isinstance(data, list):
            text = '\n'.join([json.dumps(item, ensure_ascii=False) for item in data])
        else:
            text = str(data)
        
        return {
            'source': str(path),
            'type': 'json',
            'data': data,
            'text': text,
            'size': len(text)
        }
    
    def _get_heading_level(self, style_name: str) -> int:
        """
        Determines heading level (Heading 1, 2, 3)
        """
        if 'Heading 1' in style_name or 'عنوان 1' in style_name:
            return 1
        elif 'Heading 2' in style_name or 'عنوان 2' in style_name:
            return 2
        elif 'Heading 3' in style_name or 'عنوان 3' in style_name:
            return 3
        return 0


if __name__ == "__main__":
    # Test the processor
    processor = LocalDocumentProcessor(ocr_enabled=True)
    
    # Example: Process a single file
    # doc = processor.process_file(Path("test.pdf"))
    # print(doc)
    
    # Example: Process entire directory
    # docs = processor.process_directory(Path("./documents"))
    # print(f"Processed {len(docs)} documents")
    
    print("LocalDocumentProcessor ready!")
    print(f"Supported formats: {processor.supported_formats}")
